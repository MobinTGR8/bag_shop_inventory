from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def add_inline_md(paragraph, s: str) -> None:
    """Add a string to a paragraph, supporting **bold** spans."""
    i = 0
    for m in BOLD_RE.finditer(s):
        if m.start() > i:
            paragraph.add_run(s[i : m.start()])
        r = paragraph.add_run(m.group(1))
        r.bold = True
        i = m.end()
    if i < len(s):
        paragraph.add_run(s[i:])


def is_table_row(line: str) -> bool:
    t = line.strip()
    return t.startswith("|") and t.endswith("|") and t.count("|") >= 2


def is_table_sep(line: str) -> bool:
    t = line.strip().replace(" ", "")
    if not (t.startswith("|") and t.endswith("|")):
        return False
    parts = [p for p in t.strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", p or "---") for p in parts)


def md_to_docx(md_text: str, out_path: Path) -> None:
    lines = md_text.splitlines()

    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    in_code = False
    code_lang = ""
    code_buf: list[str] = []

    table_buf: list[str] = []

    def flush_code() -> None:
        nonlocal code_buf, code_lang
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run(f"[{code_lang or 'code'} block]\n" + "\n".join(code_buf))
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        p.style = doc.styles["No Spacing"]
        code_buf = []
        code_lang = ""

    def flush_table() -> None:
        nonlocal table_buf
        if not table_buf:
            return

        rows = [r.strip() for r in table_buf if r.strip()]
        if len(rows) < 2:
            for r in rows:
                doc.add_paragraph(r)
            table_buf = []
            return

        header = [c.strip() for c in rows[0].strip("|").split("|")]
        body_rows = rows[2:] if len(rows) >= 3 else []

        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, cell_text in enumerate(header):
            hdr_cells[i].text = cell_text

        for r in body_rows:
            cols = [c.strip() for c in r.strip("|").split("|")]
            tr = table.add_row().cells
            for i in range(len(header)):
                tr[i].text = cols[i] if i < len(cols) else ""

        doc.add_paragraph("")
        table_buf = []

    for raw in lines:
        line = raw.rstrip("\n")

        # Code fences
        if line.strip().startswith("```"):
            if not in_code:
                flush_table()
                in_code = True
                code_lang = line.strip().lstrip("`").strip()
                continue
            in_code = False
            flush_code()
            continue

        if in_code:
            code_buf.append(line)
            continue

        # Tables
        if is_table_row(line):
            table_buf.append(line)
            continue
        if table_buf:
            if is_table_sep(line) or is_table_row(line):
                table_buf.append(line)
                continue
            flush_table()

        # Horizontal rules: treat as page break for a more PDF-like look
        if line.strip() == "---":
            doc.add_page_break()
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        # Blockquote
        if line.lstrip().startswith(">"):
            p = doc.add_paragraph()
            run = p.add_run(line.lstrip()[1:].lstrip())
            run.italic = True
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        # Lists
        m_num = re.match(r"^(\d+)\.\s+(.*)$", line.strip())
        if m_num:
            p = doc.add_paragraph(style="List Number")
            add_inline_md(p, m_num.group(2))
            continue

        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_md(p, line.strip()[2:])
            continue

        # Default paragraph
        p = doc.add_paragraph()
        add_inline_md(p, line)

    flush_table()
    if in_code:
        flush_code()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    if out_path.exists():
        out_path.unlink()
    doc.save(out_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: python scripts/md_to_docx.py <input.md> <output.docx>")
        return 2

    md_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])

    md_text = md_path.read_text(encoding="utf-8")
    md_to_docx(md_text, out_path)
    print(f"Wrote: {out_path.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
